"""
Genera Reporte_Entrega2.docx para el proyecto SAT-Dengue (Grupo 11).
Uso: python scripts/build_report.py
"""
import json
import os
import pickle

import numpy as np
import pandas as pd
from sklearn.linear_model import LogisticRegression
from sklearn.metrics import average_precision_score, f1_score, roc_auc_score
from sklearn.preprocessing import StandardScaler

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT    = os.path.join(os.path.dirname(__file__), "..")
FIG_DIR = os.path.join(ROOT, "data", "figures")
OUT     = os.path.join(ROOT, "Reporte_Entrega2.docx")

CALIBRI    = "Calibri"
CLR_TITLE  = RGBColor(0x16, 0x54, 0xA2)
CLR_HEAD   = RGBColor(0x1A, 0x1A, 0x1A)
BODY_SIZE  = Pt(11)
HEAD1_SIZE = Pt(14)
HEAD2_SIZE = Pt(12)

PROHIBIDAS = {
    "divipola", "municipio", "departamento", "periodo",
    "anio", "mes", "casos_grave", "casos_clasico", "brote", "es_inicio",
}
TRAIN_END = 2023
DIVIPOLAS = {"68001": "Bucaramanga", "76001": "Cali"}


# ── Helpers de formato ───────────────────────────────────────────────────────
def set_font(run, size=None, bold=False, italic=False, color=None):
    run.font.name = CALIBRI
    if size:   run.font.size = size
    run.font.bold   = bold
    run.font.italic = italic
    if color:  run.font.color.rgb = color


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    size  = HEAD1_SIZE if level == 1 else HEAD2_SIZE
    color = CLR_TITLE if level == 1 else CLR_HEAD
    set_font(run, size=size, bold=True, color=color)
    return p


def add_body(doc, text, space_after=6, italic=False, bold_label=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if bold_label:
        r = p.add_run(bold_label + " ")
        set_font(r, size=BODY_SIZE, bold=True)
    run = p.add_run(text)
    set_font(run, size=BODY_SIZE, italic=italic)
    return p


def add_figure(doc, filename, caption, width=5.5):
    path = os.path.join(FIG_DIR, filename)
    if not os.path.exists(path):
        add_body(doc, f"[Figura pendiente: {filename}]", italic=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    r = cap.add_run(caption)
    set_font(r, size=Pt(9), italic=True)


def add_table(doc, headers, rows, highlight_row=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.font.name = CALIBRI
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tcPr = hdr_cells[i]._tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1654A2")
        shd.set(qn("w:val"),  "clear")
        tcPr.append(shd)
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row_data):
            row[c_idx].text = str(val)
            for run in row[c_idx].paragraphs[0].runs:
                run.font.name = CALIBRI
                run.font.size = Pt(10)
        if highlight_row is not None and r_idx == highlight_row:
            for cell in row:
                tcPr = cell._tc.get_or_add_tcPr()
                shd  = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "E8F0FB")
                shd.set(qn("w:val"),  "clear")
                tcPr.append(shd)
    doc.add_paragraph()


# ── Datos y modelos ───────────────────────────────────────────────────────────
def load_everything():
    df = pd.read_parquet(os.path.join(ROOT, "data", "processed", "features_mensual.parquet"))
    df["divipola"] = df["divipola"].astype(str).str.zfill(5)
    feats = [c for c in df.columns
             if c not in PROHIBIDAS and pd.api.types.is_numeric_dtype(df[c])]

    train = df[df["anio"] <= TRAIN_END]
    test  = df[df["anio"] > TRAIN_END]
    val   = train[train["anio"] >= TRAIN_END - 1]
    tr    = train[train["anio"] < TRAIN_END - 1]

    with open(os.path.join(ROOT, "model", "xgb_clasico.pkl"),  "rb") as f: xgb_m = pickle.load(f)
    with open(os.path.join(ROOT, "model", "lgbm_clasico.pkl"), "rb") as f: lgb_m = pickle.load(f)
    with open(os.path.join(ROOT, "model", "gam_clasico.pkl"),  "rb") as f: gam_d = pickle.load(f)
    with open(os.path.join(ROOT, "model", "xgb_clasico_meta.json")) as f:  meta  = json.load(f)

    scaler = StandardScaler()
    X_tr_s = scaler.fit_transform(tr[feats].fillna(0))
    lr_m   = LogisticRegression(C=0.1, max_iter=1000, class_weight="balanced", random_state=42)
    lr_m.fit(X_tr_s, tr["brote"])

    X_te    = test[feats].fillna(0)
    y_te    = test["brote"]
    y_ini   = test["es_inicio"]

    gam_m     = gam_d["model"]
    gam_feats = gam_d["features"]
    gam_thr   = gam_d["best_threshold"]

    probs = {
        "XGBoost":    (xgb_m.predict_proba(X_te)[:, 1],               meta["best_threshold"]),
        "LightGBM":   (lgb_m.predict_proba(X_te)[:, 1],               0.49),
        "Logistica":  (lr_m.predict_proba(scaler.transform(X_te))[:,1], 0.63),
        "GAM":        (gam_m.predict_proba(test[gam_feats].fillna(0).values), gam_thr),
    }

    results = {}
    for name, (prob, thr) in probs.items():
        pred    = (prob >= thr).astype(int)
        auroc   = roc_auc_score(y_te, prob)
        ap      = average_precision_score(y_te, prob)
        f1      = f1_score(y_te, pred, zero_division=0)
        ini_det = int(pred[y_ini == 1].sum())
        ini_tot = int(y_ini.sum())
        results[name] = {
            "auroc": auroc, "ap": ap, "f1": f1, "thr": thr,
            "ini_det": ini_det, "ini_tot": ini_tot,
            "ini_pct": ini_det / max(ini_tot, 1) * 100,
        }

    # Baselines
    pers  = test["brote_lag_1"].fillna(0).astype(int)
    canal = (test["zona_canal_lag1"].fillna(0) >= 2).astype(int)
    for bl_name, pred_bl in [("Persistencia", pers), ("Canal endemico", canal)]:
        f1_bl  = f1_score(y_te, pred_bl, zero_division=0)
        ini_bl = int(pred_bl[y_ini == 1].sum())
        results[bl_name] = {
            "auroc": None, "ap": None, "f1": f1_bl, "thr": None,
            "ini_det": ini_bl, "ini_tot": int(y_ini.sum()),
            "ini_pct": ini_bl / max(y_ini.sum(), 1) * 100,
        }

    return meta, results, test, y_te, y_ini, feats


# ── Portada ───────────────────────────────────────────────────────────────────
def build_cover(doc):
    for _ in range(4):
        doc.add_paragraph()
    for text, size, bold, italic in [
        ("SAT-Dengue",                                          Pt(26), True,  False),
        ("Sistema de Alerta Temprana para Brotes de Dengue en Colombia",
                                                               Pt(14), False, True),
        ("",                                                    Pt(12), False, False),
        ("Entrega 2",                                           Pt(14), True,  False),
        ("Proyecto MAIA — Procesamiento y Despliegue de Soluciones",
                                                               Pt(12), False, False),
        ("",                                                    Pt(11), False, False),
        ("Grupo 11",                                            Pt(12), True,  False),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_font(r, size=size, bold=bold, italic=italic,
                 color=CLR_TITLE if "SAT" in text or "Entrega" in text else None)
    for name in ["Stevan Castro", "Juan Mauricio Ramirez", "Nilara Vega"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(name)
        set_font(r, size=Pt(11))
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Septiembre de 2026")
    set_font(r, size=Pt(11))
    doc.add_page_break()


# ── Seccion 1 ─────────────────────────────────────────────────────────────────
def build_s1(doc):
    add_heading(doc, "1. Resumen del problema")

    add_heading(doc, "1.1 Contexto y pregunta de negocio", level=2)
    add_body(doc,
        "El dengue es la arbovirosis de mayor expansión global, con estimaciones de 390 millones "
        "de infecciones anuales. En Colombia, el Sistema de Vigilancia en Salud Pública (SIVIGILA) "
        "reporta entre 30.000 y 170.000 casos de dengue por año, con picos pronunciados en "
        "municipios urbanos de las regiones Andina y Caribe. Las secretarías de salud municipales "
        "son las responsables de activar medidas de respuesta (fumigación, bloqueo focal, refuerzo "
        "hospitalario) ante la llegada de un brote, pero la herramienta de vigilancia actual, el "
        "canal endémico, solo emite señal de alerta una vez el municipio ya lleva al menos un mes "
        "en zona de exceso. Para entonces, la transmisión está en su pico y la capacidad de "
        "respuesta preventiva es muy limitada."
    )
    add_body(doc,
        "La pregunta de negocio que guía el proyecto es: ¿Es posible predecir con un mes de "
        "anticipación si un municipio colombiano superará el umbral P75 del canal endémico "
        "histórico de dengue total, permitiendo activar alertas antes de que el brote se "
        "consolide?"
    )

    add_heading(doc, "1.2 Objetivo y alcance", level=2)
    add_body(doc,
        "El objetivo es desarrollar un sistema de clasificación mensual por municipio que "
        "distinga los meses que entrarán en zona de exceso epidémico (brote) de los que no, "
        "utilizando únicamente información disponible al momento de la predicción: casos "
        "acumulados de meses anteriores, indicadores derivados del canal endémico y covariables "
        "de estacionalidad. El sistema piloto se evalúa en Bucaramanga (código DIVIPOLA 68001) "
        "y Cali (76001), municipios seleccionados por su diferente perfil epidémico y volumen "
        "de casos. Los datos se agregan por fecha de inicio de síntomas (campo INI_SIN de "
        "SIVIGILA), que es la referencia estándar del canal endémico."
    )

    add_heading(doc, "1.3 Conjuntos de datos", level=2)
    add_body(doc,
        "La fuente principal es SIVIGILA 2007-2025. Se utilizaron los registros de dengue total "
        "(casos notificados de dengue, incluyendo dengue con signos de alarma) y dengue grave "
        "como series separadas. El dataset procesado contiene 253.992 filas (municipio-mes) "
        "con 38 columnas para 1.114 municipios, construido a partir de un panel completo que "
        "incluye ceros para los municipio-mes sin notificaciones. Las fechas se asignan al mes "
        "de inicio de síntomas, con fecha de notificación como respaldo cuando aquella está "
        "ausente. Se prevé incorporar en fases futuras datos climáticos mensuales de temperatura "
        "(ERA5) y precipitación (CHIRPS) a través de Google Earth Engine, actualmente en proceso "
        "de descarga."
    )

    add_heading(doc, "1.4 Cambios respecto a Entrega 1", level=2)
    add_body(doc,
        "La decisión de cambio más relevante fue la sustitución de la variable objetivo. En "
        "la Entrega 1 el modelo intentaba predecir brotes de dengue grave (casos_grave > P75). "
        "El análisis de los registros históricos reveló que la reclasificación de la OMS en 2009 "
        "modificó drásticamente los criterios diagnósticos de dengue grave, produciendo una "
        "caída de aproximadamente 45 veces en los reportes de Bucaramanga (de un promedio de "
        "1.752 casos/año antes de 2009 a entre 1 y 18 casos/año después de 2011). Esta "
        "discontinuidad hace que el canal endémico de dengue grave sea epidemiológicamente "
        "ininterpretable y que la tasa de positivos en los folds de validación sea inferior "
        "al 4%, impidiendo cualquier modelo con valor predictivo."
    )
    add_body(doc,
        "El equipo adoptó dengue total (casos_clasico) como objetivo, lo cual eleva la tasa "
        "de positivos al 14,1% en validación, produce umbrales P25/P75 interpretables "
        "(Bucaramanga: 158-378 casos/mes, Cali: 480-1.580 casos/mes) y permite cuantificar "
        "el valor añadido del modelo frente a la práctica actual. Adicionalmente se corrigió "
        "un error en el cálculo de promedios móviles que no agrupaba por municipio, y se "
        "eliminó la filtración de datos futuros en el cálculo del canal endémico dentro de "
        "los folds de validación cruzada."
    )


# ── Seccion 2 ─────────────────────────────────────────────────────────────────
def build_s2(doc, meta, results):
    add_heading(doc, "2. Modelos desarrollados y su evaluación")

    add_heading(doc, "2.1 Pipeline de características", level=2)
    add_body(doc,
        "El pipeline parte de los archivos SIVIGILA crudos y produce el panel de features "
        "en dos etapas. La primera (build_panel.py) consolida los registros diarios en "
        "conteos mensuales por municipio, completa el panel con ceros para municipio-mes sin "
        "casos, y excluye registros con lugar de ocurrencia exterior o desconocido. La segunda "
        "etapa (build_features.py) calcula 28 predictoras: lags de 1, 2, 3, 4 y 6 meses y "
        "promedios móviles de tres meses para casos_clasico y casos_grave; lags de temperatura "
        "y precipitación (actualmente nulos); componentes trigonométricas de estacionalidad "
        "(mes_sin, mes_cos); y variables del canal endémico (p25, p75, zona_canal_lag1, "
        "es_endemico, brote_lag_1). La variable objetivo es brote, definida como "
        "casos_clasico > p75 en ese mes."
    )
    add_body(doc,
        "La variable es_inicio marca el primer mes de un episodio de brote, es decir, el "
        "mes en que el municipio transiciona de zona normal o de alerta a zona de exceso. "
        "Esta distinción es central para evaluar el valor real del modelo: un sistema de "
        "alerta temprana debe detectar inicios, no solo confirmar que un brote ya está "
        "activo."
    )

    add_heading(doc, "2.2 Canal endémico y la distinción inicio-continuacion", level=2)
    add_body(doc,
        "El canal endémico es la herramienta estándar de vigilancia epidemiológica colombiana. "
        "Para cada municipio y mes del año se calculan el P25 y el P75 de casos históricos "
        "(período de referencia 2007-2022). Un mes se clasifica en zona normal si los casos "
        "están por debajo del P25, en zona de alerta si se ubican entre P25 y P75, y en zona "
        "de exceso o brote si superan el P75."
    )
    add_body(doc,
        "La práctica actual de las secretarías de salud equivale al baseline de persistencia: "
        "se emite alerta cuando el mes anterior ya estaba en zona de exceso. Esto significa "
        "que el sistema actual solo detecta la continuación de un brote, nunca su inicio. "
        "En términos epidemiológicos: un brote es un estado, no un evento puntual. El valor "
        "real de un modelo predictivo está en detectar la transición hacia ese estado, no "
        "en confirmar que el estado ya existe."
    )
    add_figure(doc, "report_canal_endemico.png",
               "Figura 1. Canal endémico histórico (P25/P75) y casos de dengue total mensual "
               "para Bucaramanga y Cali (2007-2025). Las barras se colorean por zona: "
               "verde (normal), naranja (alerta), rojo (exceso/brote).", width=5.8)
    add_figure(doc, "report_estacionalidad_brotes.png",
               "Figura 2. Porcentaje de meses en brote por mes del año calendario. "
               "Se observa estacionalidad marcada con picos en el primer y tercer trimestre.",
               width=5.0)

    add_heading(doc, "2.3 Modelos entrenados", level=2)
    add_body(doc,
        "Se entrenaron cinco modelos sobre la misma partición temporal (entrenamiento: 2007-2023, "
        "prueba: 2024-2025), con validación interna sobre 2022-2023 para selección de umbral. "
        "La variable objetivo en todos los casos es brote (casos_clasico > P75)."
    )
    add_body(doc, text=(
        "XGBoost: clasificador de gradient boosting con 400 estimadores, profundidad máxima "
        f"6, learning rate 0,05 y scale_pos_weight 5,42 para el desbalance de clases. "
        f"Early stopping en 30 rondas sobre AUCPR determinó {meta['best_iteration']} iteraciones "
        f"óptimas. Umbral calibrado en {meta['best_threshold']} maximizando F1 en validación."),
        bold_label="XGBoost.")
    add_body(doc, text=(
        "Gradient boosting implementado con LightGBM, arquitectura basada en histogramas con "
        "63 hojas por árbol y early stopping a 30 rondas. Significativamente más rápido que "
        "XGBoost para el mismo volumen de datos. Umbral óptimo 0,49."),
        bold_label="LightGBM.")
    add_body(doc, text=(
        "Modelo de ensemble por bagging con 500 árboles, profundidad máxima 12 y mínimo de "
        "20 muestras por hoja. Sirve como referencia de bagging frente al boosting de los "
        "modelos anteriores. Umbral óptimo 0,60."),
        bold_label="Random Forest.")
    add_body(doc, text=(
        "Referencia lineal con regularización L2 (C=0,1) y pesos balanceados por clase. "
        "Permite cuantificar el aporte de la no-linealidad del modelo ganador. Umbral 0,63."),
        bold_label="Regresion Logistica.")
    add_body(doc, text=(
        "GAM logístico (LogisticGAM de pygam) con términos de spline cúbico para variables "
        "de tendencia y estacionalidad, y términos lineales para variables binarias y "
        "categóricas. El parámetro lambda se seleccionó por búsqueda en grilla. Ofrece "
        "interpretabilidad de efectos parciales por variable."),
        bold_label="GAM.")

    add_heading(doc, "2.4 Gestion de experimentos con MLflow", level=2)
    add_body(doc,
        "Todos los entrenamientos se registran en el experimento 'dengue-brote-clasico' de "
        "MLflow. El servidor está desplegado en una instancia EC2 de AWS, accesible para "
        "todos los integrantes del equipo mediante la variable de entorno MLFLOW_TRACKING_URI. "
        "Para cada corrida se registran hiperparámetros, métricas de validación y prueba, "
        "umbral óptimo y el modelo serializado como artefacto. Los modelos XGBoost y LightGBM "
        "se registran además en el Registro de Modelos de MLflow, lo que permite cargarlos "
        "desde el tablero de predicciones sin acceso al sistema de archivos local. "
        "El repositorio incluye un archivo MLproject que permite lanzar cualquier entrenamiento "
        "con: mlflow run . -e train_xgboost (o train_logistic, train_gam, train_rf, train_lgbm)."
    )

    add_heading(doc, "2.5 Evaluacion comparativa", level=2)
    add_body(doc,
        "La Tabla 1 resume el desempeño de todos los modelos sobre el conjunto de prueba "
        "2024-2025, que comprende 26.736 municipio-mes con una tasa de brote del 43%, "
        "correspondiente a un período epidémico severo de dengue en Colombia."
    )

    headers = ["Modelo", "AUROC", "AP", "F1", "Umbral", "Inicios det."]
    rows = []
    model_order = ["XGBoost", "LightGBM", "Logistica", "GAM",
                   "Persistencia", "Canal endemico"]
    auroc_fmt = {
        "XGBoost":   f"{results['XGBoost']['auroc']:.4f}",
        "LightGBM":  f"{results['LightGBM']['auroc']:.4f}",
        "Logistica": f"{results['Logistica']['auroc']:.4f}",
        "GAM":       f"{results['GAM']['auroc']:.4f}",
        "Persistencia":    "—",
        "Canal endemico":  "—",
    }
    for name in model_order:
        r = results[name]
        ini = f"{r['ini_det']}/{r['ini_tot']} ({r['ini_pct']:.0f}%)"
        thr = f"{r['thr']:.2f}" if r['thr'] else "—"
        ap  = f"{r['ap']:.4f}"  if r['ap']  else "—"
        rows.append([name, auroc_fmt[name], ap, f"{r['f1']:.3f}", thr, ini])

    add_table(doc, headers, rows, highlight_row=0)
    add_body(doc,
        "Tabla 1. Comparativa de modelos en el conjunto de prueba 2024-2025. "
        "La columna 'Inicios det.' indica cuántos de los 2.461 meses de inicio de brote "
        "(es_inicio=1) cada modelo identifica correctamente con su umbral óptimo.",
        italic=True, space_after=10)

    add_figure(doc, "report_roc_pr.png",
               "Figura 3. Curvas ROC (izq.) y Precisión-Recall (der.) para XGBoost y "
               "Regresión Logística sobre el conjunto de prueba 2024-2025.", width=5.8)

    add_body(doc,
        "La Figura 4 muestra el hallazgo más relevante desde la perspectiva del negocio: "
        "el porcentaje de inicios de brote detectados por cada modelo. La persistencia "
        "(práctica actual) detecta 0% de los inicios, ya que por definición solo puede "
        "confirmar un brote en curso, nunca anticiparlo. El canal endémico detecta el 18%. "
        "XGBoost y LightGBM alcanzan el 32% y 28% respectivamente, casi el doble que la "
        "práctica actual. Este es el valor agregado real del modelo: anticipar el inicio "
        "de un episodio epidémico con al menos un mes de antelación."
    )
    add_figure(doc, "report_inicios_brote.png",
               "Figura 4. Porcentaje de inicios de brote (es_inicio=1) detectados por cada "
               "modelo con su umbral óptimo sobre el conjunto de prueba 2024-2025. "
               "La línea naranja marca el 18% de la práctica actual (canal endémico).",
               width=5.5)

    add_figure(doc, "report_feature_importance.png",
               "Figura 5. Importancia de features (gain) del modelo XGBoost. Las variables "
               "de tendencia reciente dominan, seguidas por el estado del canal endémico "
               "y la estacionalidad.", width=4.8)


# ── Seccion 3 ─────────────────────────────────────────────────────────────────
def build_s3(doc):
    add_heading(doc, "3. Tablero de predicciones")
    add_body(doc,
        "El tablero de predicciones fue desarrollado con el framework Plotly Dash y está "
        "orientado a los equipos de vigilancia de las secretarías de salud de Bucaramanga "
        "y Cali. Presenta la serie histórica de dengue total desde 2007 con los límites del "
        "canal endémico superpuestos, la probabilidad de brote estimada por XGBoost para "
        "cada mes histórico y una predicción para el mes siguiente al último dato disponible, "
        "presentada como un medidor de riesgo con clasificación de zona (normal, alerta o exceso)."
    )
    add_body(doc,
        "El tablero incluye tres pestañas (Bucaramanga, Cali y comparación entre ciudades) "
        "y cuatro indicadores clave por ciudad: total de meses analizados, meses en brote, "
        "inicios de brote históricos y pico máximo de casos. El sistema está preparado para "
        "cargar el modelo directamente desde el servidor MLflow configurando la variable de "
        "entorno MLFLOW_MODEL_URI, lo que permite actualizar el modelo en producción sin "
        "redesplegar el tablero. Para el despliegue en EC2 se configura DASH_HOST=0.0.0.0 "
        "y el puerto correspondiente."
    )


# ── Seccion 4 ─────────────────────────────────────────────────────────────────
def build_s4(doc, meta, results):
    add_heading(doc, "4. Observaciones y conclusiones")
    add_body(doc,
        "El cambio de objetivo de dengue grave a dengue total fue la decisión técnica más "
        "relevante de esta entrega y la que desbloqueó el resto del trabajo. Sin ella, "
        "cualquier modelo habría operado sobre una señal rota por el cambio de clasificación "
        "de la OMS, produciendo resultados sin validez epidemiológica."
    )
    xgb = results["XGBoost"]
    lgb = results["LightGBM"]
    add_body(doc,
        f"XGBoost y LightGBM son los modelos con mayor desempeño (AUROC {xgb['auroc']:.4f} "
        f"y {lgb['auroc']:.4f} respectivamente), con métricas prácticamente idénticas. "
        "La paridad entre ambos sugiere que el límite de desempeño actual no es arquitectural "
        "sino de información: las 28 features disponibles, en particular los lags de casos "
        "y las variables del canal endémico, capturan la mayor parte de la señal predictiva "
        "accesible sin datos climáticos."
    )
    pers = results["Persistencia"]
    xgb_ini_pct = xgb["ini_pct"]
    add_body(doc,
        f"El hallazgo más importante para el caso de uso es la detección de inicios de brote. "
        f"La práctica actual (persistencia) detecta 0% de los {xgb['ini_tot']:,} meses de "
        f"inicio de brote en el período de prueba. XGBoost detecta el {xgb_ini_pct:.0f}%, "
        "casi el doble que el canal endémico. Esto cuantifica el valor operativo del modelo: "
        "si una secretaría de salud adopta el sistema, podría anticipar aproximadamente un "
        "tercio de los episodios epidémicos antes de que se consoliden, frente a ninguno con "
        "la herramienta actual."
    )
    add_body(doc,
        "Las principales limitaciones son la ausencia de datos climáticos, que podrían "
        "extender el horizonte de anticipación; la cobertura limitada al piloto de dos "
        "municipios; y la granularidad mensual, que impide capturar dinámicas intra-mensuales "
        "relevantes para la activación de respuestas inmediatas. El período de prueba "
        "2024-2025 presentó una tasa de brote del 43%, notablemente superior al histórico "
        "(14-23%), lo que indica un año epidémico severo en Colombia y puede sesgar al alza "
        "las métricas de prueba frente a años más típicos."
    )


# ── Seccion 5 ─────────────────────────────────────────────────────────────────
def build_s5(doc):
    add_heading(doc, "5. Reporte de trabajo en equipo")
    add_body(doc,
        "El trabajo de la Entrega 2 se distribuyó de la siguiente forma entre los integrantes "
        "del Grupo 11:"
    )
    roles = [
        ("Stevan Castro",
         "Análisis del cambio de clasificación OMS y decisión de cambio de objetivo; "
         "implementación del pipeline de características (build_panel.py, build_features.py); "
         "corrección del bug de promedios móviles y la filtración de datos futuros en CV; "
         "análisis de baselines incluyendo es_inicio; configuración del servidor MLflow en "
         "EC2 y registro de todos los experimentos del equipo."),
        ("Juan Mauricio Ramirez",
         "Diseño e implementación del tablero de predicciones en Plotly Dash; "
         "integración con el servidor MLflow para carga dinámica del modelo desde el "
         "registro de modelos; despliegue del tablero en la instancia EC2."),
        ("Nilara Vega",
         "Entrenamiento y evaluación de los cinco modelos (XGBoost, LightGBM, Random Forest, "
         "Regresión Logística, GAM); evaluación específica de detección de inicios de brote "
         "(es_inicio); generación de figuras de análisis; redacción del reporte; integración "
         "de todos los scripts de entrenamiento con MLflow (MLproject, registro de modelos)."),
    ]
    for nombre, contrib in roles:
        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(5)
        p.paragraph_format.space_before = Pt(2)
        r1 = p.add_run(f"{nombre}: ")
        set_font(r1, size=BODY_SIZE, bold=True)
        r2 = p.add_run(contrib)
        set_font(r2, size=BODY_SIZE)


# ── Main ──────────────────────────────────────────────────────────────────────
def main():
    print("Cargando datos y modelos...")
    meta, results, test, y_te, y_ini, feats = load_everything()

    print("Construyendo documento...")
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.18)
        section.right_margin  = Inches(1.18)
    style = doc.styles["Normal"]
    style.font.name = CALIBRI
    style.font.size = BODY_SIZE

    build_cover(doc)
    build_s1(doc)
    build_s2(doc, meta, results)
    doc.add_page_break()
    build_s3(doc)
    build_s4(doc, meta, results)
    build_s5(doc)

    doc.save(OUT)
    print(f"Reporte guardado: {OUT}")


if __name__ == "__main__":
    main()
