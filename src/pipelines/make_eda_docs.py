"""
Genera dos documentos Word con los resultados del EDA:
  - EDA_Dengue_Clasico.docx
  - EDA_Dengue_Grave.docx
"""

import os
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE   = r"C:\Users\nilara\Downloads\dengue_project"
FIGS   = os.path.join(BASE, "data", "figures")
PROC   = os.path.join(BASE, "data", "processed")
OUT    = r"C:\Users\nilara\Downloads"

CSV_210      = os.path.join(PROC, "sivigila_dengue_consolidado.csv")
CSV_220      = os.path.join(PROC, "sivigila_dengue_grave_consolidado.csv")
CSV_ENDEMICOS = os.path.join(PROC, "municipios_endemicos.csv")

EPIDEMIC_YEARS = [2010, 2013, 2016, 2019, 2023, 2024]


# ── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_figure(doc, path, caption, width=5.8):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].italic = True
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = RGBColor(0x5C, 0x6B, 0x84)
    else:
        doc.add_paragraph(f"[Figura no encontrada: {path}]")

def stats_table(doc, df_stat, col_widths=None):
    t = doc.add_table(rows=1, cols=len(df_stat.columns))
    t.style = 'Table Grid'
    hdr = t.rows[0].cells
    for i, col in enumerate(df_stat.columns):
        hdr[i].text = str(col)
        hdr[i].paragraphs[0].runs[0].bold = True
        set_cell_bg(hdr[i], '0E1E3C')
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for _, row in df_stat.iterrows():
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph()


# ── Calcular estadísticas ────────────────────────────────────────────────────

print("Cargando datos para estadísticas...")
df_c = pd.read_csv(CSV_210, usecols=['source_file','ANO','SEMANA',
                                      'Municipio_ocurrencia','Departamento_ocurrencia',
                                      'TIP_CAS','CON_FIN','PAC_HOS','SEXO','EDAD'],
                   dtype=str, low_memory=False)
df_c['ANO']    = pd.to_numeric(df_c['ANO'],    errors='coerce')
df_c['SEMANA'] = pd.to_numeric(df_c['SEMANA'], errors='coerce')
df_c['source_file'] = pd.to_numeric(df_c['source_file'], errors='coerce')

df_g = pd.read_csv(CSV_220, usecols=['source_file','ANO','SEMANA',
                                      'Municipio_ocurrencia','Departamento_ocurrencia',
                                      'TIP_CAS','CON_FIN','PAC_HOS','SEXO','EDAD'],
                   dtype=str, low_memory=False)
df_g['ANO']    = pd.to_numeric(df_g['ANO'],    errors='coerce')
df_g['SEMANA'] = pd.to_numeric(df_g['SEMANA'], errors='coerce')
df_g['source_file'] = pd.to_numeric(df_g['source_file'], errors='coerce')

endemicos = pd.read_csv(CSV_ENDEMICOS)

# Por año — clásico
por_ano_c = df_c.groupby('source_file').size().reset_index(name='Registros')
por_ano_c.columns = ['Año', 'Registros']
por_ano_c['Epidemia'] = por_ano_c['Año'].apply(lambda y: 'Si' if y in EPIDEMIC_YEARS else 'No')
por_ano_c['Registros'] = por_ano_c['Registros'].apply(lambda x: f"{x:,}")

# Por año — grave
por_ano_g = df_g.groupby('source_file').size().reset_index(name='Registros')
por_ano_g.columns = ['Año', 'Registros']
por_ano_g['Epidemia'] = por_ano_g['Año'].apply(lambda y: 'Si' if y in EPIDEMIC_YEARS else 'No')
por_ano_g['Registros'] = por_ano_g['Registros'].apply(lambda x: f"{x:,}")

# Tasa de gravedad
clasico_ano = df_c.groupby('source_file').size().rename('clasico')
grave_ano   = df_g.groupby('source_file').size().rename('grave')
tasa = pd.concat([clasico_ano, grave_ano], axis=1).dropna()
tasa['Tasa (%)'] = (tasa['grave'] / tasa['clasico'] * 100).round(2)
tasa = tasa.reset_index().rename(columns={'source_file': 'Año', 'clasico': 'Clasico', 'grave': 'Grave'})
tasa['Clasico'] = tasa['Clasico'].apply(lambda x: f"{int(x):,}")
tasa['Grave']   = tasa['Grave'].apply(lambda x: f"{int(x):,}")
tasa['Tasa (%)']= tasa['Tasa (%)'].apply(lambda x: f"{x:.2f}%")
tasa['Año']     = tasa['Año'].astype(int)

# Top 10 municipios — clásico
top_mun_c = (df_c.groupby(['Municipio_ocurrencia','Departamento_ocurrencia'])
               .size().reset_index(name='Casos')
               .sort_values('Casos', ascending=False).head(10))
top_mun_c['Casos'] = top_mun_c['Casos'].apply(lambda x: f"{x:,}")
top_mun_c.columns = ['Municipio', 'Departamento', 'Casos']
top_mun_c = top_mun_c.reset_index(drop=True)

# Top 10 municipios — grave
top_mun_g = (df_g.groupby(['Municipio_ocurrencia','Departamento_ocurrencia'])
               .size().reset_index(name='Casos')
               .sort_values('Casos', ascending=False).head(10))
top_mun_g['Casos'] = top_mun_g['Casos'].apply(lambda x: f"{x:,}")
top_mun_g.columns = ['Municipio', 'Departamento', 'Casos']
top_mun_g = top_mun_g.reset_index(drop=True)

# Hospitalización — grave
hosp_g = df_g['PAC_HOS'].value_counts().rename({1:'Si',2:'No','1':'Si','2':'No'})
pct_hosp = (hosp_g / hosp_g.sum() * 100).round(1)

# Mortalidad — grave
mort_g = df_g['CON_FIN'].value_counts().rename({'1':'Vivo','2':'Muerto','3':'No sabe',
                                                  1:'Vivo',2:'Muerto',3:'No sabe'})
pct_mort = (mort_g / mort_g.sum() * 100).round(1)

print("Estadísticas calculadas.")


# ══ DOCUMENTO 1: DENGUE CLASICO ══════════════════════════════════════════════

print("Creando EDA_Dengue_Clasico.docx...")
doc = Document()

# Márgenes
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(1.2)

# Portada
heading(doc, 'Exploración de Datos — Dengue Clásico (COD_EVE 210)', 1)
body(doc, 'Sistema de Alerta Temprana de Dengue · Grupo 11 · MAIA PDS · Universidad de los Andes')
body(doc, 'Fuente: SIVIGILA / INS · Período: 2007–2024 · Evento 210 (dengue sin/con signos de alarma)')
doc.add_paragraph()

# Resumen general
heading(doc, '1. Resumen general', 2)
resumen_data = pd.DataFrame({
    'Parámetro': [
        'Total de registros', 'Años cubiertos', 'Columnas por registro',
        'Municipios con al menos 1 caso', 'Municipios endémicos (D4)',
        'Años de epidemia mayor', 'Período de referencia corredor endémico'
    ],
    'Valor': [
        '1,585,040', '2007–2024 (18 años)', '69',
        f"{df_c['Municipio_ocurrencia'].nunique():,}",
        f"{len(endemicos):,}",
        '2010, 2013, 2016, 2019, 2023, 2024',
        '2007–2022 (excluyendo epidémicos)'
    ]
})
stats_table(doc, resumen_data)

body(doc, (
    'El dataset de dengue clásico contiene 1,585,040 notificaciones individuales distribuidas '
    'en 18 años (2007–2024). Cada fila corresponde a un caso notificado al sistema SIVIGILA '
    'con 69 variables clínicas, epidemiológicas y geográficas. La estructura es idéntica en '
    'todos los archivos anuales, lo que permite consolidarlos sin transformaciones adicionales.'
))

# Serie temporal
heading(doc, '2. Serie temporal por año', 2)
body(doc, (
    'Los años epidémicos (2010, 2013, 2016, 2019, 2023, 2024) concentran la mayor carga de '
    'casos. El año 2024 es el de mayor registro histórico con 309,627 notificaciones, '
    'más del doble que cualquier año no epidémico anterior.'
))
add_figure(doc, os.path.join(FIGS, '01_clasico_por_ano.png'),
           'Figura 1. Casos de dengue clásico por año (2007–2024). En rojo: años epidémicos.')
stats_table(doc, por_ano_c)

# Corredor endémico
heading(doc, '3. Corredor endémico nacional', 2)
body(doc, (
    'El corredor endémico se construye sobre el período de referencia 2007–2022, '
    'excluyendo los años de epidemia mayor. Para cada semana epidemiológica se calculan '
    'los percentiles P25 (zona de éxito), P50 (mediana, zona endémica) y P75 '
    '(umbral de alerta/epidemia), siguiendo la metodología de Bortman (1999) adoptada '
    'por el INS Colombia.'
))
add_figure(doc, os.path.join(FIGS, '02_corredor_endemico_nacional.png'),
           'Figura 2. Corredor endémico nacional. La línea roja discontinua (P75) es el umbral de alerta epidémica.')

# Geografía
heading(doc, '4. Distribución geográfica', 2)
body(doc, (
    f'Se identificaron {df_c["Municipio_ocurrencia"].nunique():,} municipios con al menos '
    f'un caso notificado en el período. Aplicando el criterio D4 del proyecto '
    f'(≥10 años con casos y ≥200 casos acumulados), {len(endemicos):,} municipios '
    f'califican como endémicos y conforman el universo del modelo predictivo.'
))
body(doc, 'Top 10 municipios por total de casos (2007–2024):')
stats_table(doc, top_mun_c)

# Calidad
heading(doc, '5. Calidad de los datos', 2)
body(doc, (
    'Los 18 archivos anuales comparten exactamente las mismas 69 columnas sin diferencias '
    'estructurales. No se detectaron registros duplicados por CONSECUTIVE entre el dataset '
    'de dengue clásico y el de dengue grave. Las variables con mayor porcentaje de valores '
    'nulos corresponden a campos de seguimiento clínico que no siempre se diligencian en '
    'la notificación inicial (p. ej., condición final, fecha de hospitalización).'
))

# Conclusiones
heading(doc, '6. Conclusiones para el modelado', 2)
body(doc, (
    f'El dataset es suficientemente amplio (1,585,040 registros) y consistente para '
    f'construir modelos de predicción. Los {len(endemicos):,} municipios endémicos '
    f'identificados tienen series históricas de al menos 10 años, lo que garantiza '
    f'corredores endémicos estadísticamente robustos. El horizonte de pronóstico de '
    f'4 a 8 semanas es viable dado el patrón estacional observado en la semana '
    f'epidemiológica, con picos consistentes en las semanas 5–12 y 35–45.'
))

path_c = os.path.join(OUT, 'EDA_Dengue_Clasico.docx')
doc.save(path_c)
print(f"Guardado: {path_c}")


# ══ DOCUMENTO 2: DENGUE GRAVE ════════════════════════════════════════════════

print("Creando EDA_Dengue_Grave.docx...")
doc2 = Document()

for section in doc2.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(1.2)

heading(doc2, 'Exploración de Datos — Dengue Grave (COD_EVE 220)', 1)
body(doc2, 'Sistema de Alerta Temprana de Dengue · Grupo 11 · MAIA PDS · Universidad de los Andes')
body(doc2, 'Fuente: SIVIGILA / INS · Período: 2007–2024 · Evento 220 (dengue grave)')
doc2.add_paragraph()

# Resumen
heading(doc2, '1. Resumen general', 2)
resumen_g = pd.DataFrame({
    'Parámetro': [
        'Total de registros', 'Años cubiertos', 'Columnas por registro',
        'Municipios con al menos 1 caso grave',
        'Tasa de gravedad promedio (grave/clasico)',
        'Relación con dataset clásico'
    ],
    'Valor': [
        '48,823', '2007–2024 (18 años)', '69',
        f"{df_g['Municipio_ocurrencia'].nunique():,}",
        f"{(len(df_g) / len(df_c) * 100):.2f}%",
        'Independiente — 0 CONSECUTIVE duplicados'
    ]
})
stats_table(doc2, resumen_g)
body(doc2, (
    'El dataset de dengue grave contiene 48,823 notificaciones individuales de casos '
    'que cumplen los criterios clínicos de gravedad (dengue con signos de alarma severos, '
    'choque, afectación orgánica grave). Comparte las mismas 69 columnas que el dataset '
    'de dengue clásico y no presenta registros duplicados entre ambos conjuntos, '
    'confirmando que son notificaciones independientes bajo códigos de evento distintos.'
))

# Serie temporal
heading(doc2, '2. Serie temporal por año', 2)
body(doc2, (
    'Los años epidémicos concentran también los picos de dengue grave, con 2010 como '
    'el año de mayor carga histórica (9,753 casos). A partir de 2011 se observa una '
    'reducción sostenida, posiblemente asociada a mejoras en el protocolo de atención '
    'temprana y la reclasificación clínica de casos.'
))
add_figure(doc2, os.path.join(FIGS, '03_grave_por_ano.png'),
           'Figura 3. Casos de dengue grave por año (2007–2024). En rojo: años epidémicos.')
stats_table(doc2, por_ano_g)

# Tasa de gravedad
heading(doc2, '3. Tasa de gravedad respecto al dengue clásico', 2)
body(doc2, (
    'La tasa de gravedad se define como el cociente entre casos graves (220) y casos '
    'clásicos (210) por año. Esta métrica es relevante para el modelo porque indica '
    'qué proporción de la carga epidémica se convierte en casos críticos que demandan '
    'hospitalización y recursos especializados.'
))
add_figure(doc2, os.path.join(FIGS, '04_tasa_gravedad.png'),
           'Figura 4. Tasa de gravedad anual (casos graves / casos clásicos). La línea discontinua indica el promedio histórico.')
stats_table(doc2, tasa)

# Geografía
heading(doc2, '4. Distribución geográfica', 2)
body(doc2, (
    f'Los {df_g["Municipio_ocurrencia"].nunique():,} municipios con casos graves '
    f'son un subconjunto de los municipios con dengue clásico. La concentración '
    f'geográfica es similar, con los valles interandinos y la Costa Caribe liderando '
    f'la carga de casos graves.'
))
body(doc2, 'Top 10 municipios por total de casos graves (2007–2024):')
stats_table(doc2, top_mun_g)

# Severidad
heading(doc2, '5. Indicadores de severidad', 2)

hosp_df = pd.DataFrame({
    'Hospitalizado (PAC_HOS)': ['Si (1)', 'No (2)'],
    'Registros': [
        f"{df_g[df_g['PAC_HOS'].isin(['1',1])].shape[0]:,}",
        f"{df_g[df_g['PAC_HOS'].isin(['2',2])].shape[0]:,}"
    ]
})
stats_table(doc2, hosp_df)

mort_df = pd.DataFrame({
    'Desenlace (CON_FIN)': ['Vivo (1)', 'Muerto (2)', 'No sabe (3)'],
    'Registros': [
        f"{df_g[df_g['CON_FIN'].isin(['1',1])].shape[0]:,}",
        f"{df_g[df_g['CON_FIN'].isin(['2',2])].shape[0]:,}",
        f"{df_g[df_g['CON_FIN'].isin(['3',3])].shape[0]:,}"
    ]
})
stats_table(doc2, mort_df)

# Calidad
heading(doc2, '6. Calidad de los datos', 2)
body(doc2, (
    'Estructura idéntica a los archivos de dengue clásico (69 columnas). '
    'No hay registros duplicados entre ambos datasets (verificado por CONSECUTIVE). '
    'Los archivos de 2025 fueron excluidos del análisis por estar fuera del período '
    'de referencia del proyecto (2007–2024).'
))

# Conclusiones
heading(doc2, '7. Conclusiones para el modelado', 2)
body(doc2, (
    'El dataset de dengue grave aporta una dimensión crítica para el sistema de alerta '
    'temprana: permite estimar no solo el riesgo de exceso de casos, sino también '
    'la probabilidad de que esos casos sean graves, lo que tiene implicaciones directas '
    'en la planificación de recursos hospitalarios. La tasa de gravedad histórica '
    f'promedio de {(len(df_g)/len(df_c)*100):.1f}% puede usarse como variable '
    'objetivo complementaria o como feature del modelo principal.'
))

path_g = os.path.join(OUT, 'EDA_Dengue_Grave.docx')
doc2.save(path_g)
print(f"Guardado: {path_g}")
print("\nListo.")
