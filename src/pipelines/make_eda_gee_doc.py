import os, pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FIGS  = r"C:\Users\nilara\Downloads\dengue_project\data\figures"
CSV   = r"C:\Users\nilara\Downloads\dengue_project\data\processed\gee_lst_municipios.csv"
OUT   = r"C:\Users\nilara\Downloads\EDA_GEE_LST.docx"

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_figure(doc, path, caption, width=5.8):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].italic = True
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = RGBColor(0x5C, 0x6B, 0x84)

def stats_table(doc, df_stat):
    t = doc.add_table(rows=1, cols=len(df_stat.columns))
    t.style = 'Table Grid'
    hdr = t.rows[0].cells
    for i, col in enumerate(df_stat.columns):
        hdr[i].text = str(col)
        hdr[i].paragraphs[0].runs[0].bold = True
        set_cell_bg(hdr[i], '0E1E3C')
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for _, row in df_stat.iterrows():
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph()

# Cargar datos para estadisticas (solo columnas necesarias)
print("Cargando datos...")
df = pd.read_csv(CSV, usecols=['source_file','ADM1_NAME','ADM2_CODE','ADM2_NAME',
                                'fecha','lst_celsius'], low_memory=False,
                 parse_dates=['fecha'])
df_v = df.dropna(subset=['lst_celsius'])

nulos_ano = df.groupby('source_file')['lst_celsius'].apply(
    lambda x: x.isna().mean()*100).reset_index()
nulos_ano.columns = ['Año', 'Nulos (%)']
nulos_ano['Nulos (%)'] = nulos_ano['Nulos (%)'].apply(lambda x: f"{x:.1f}%")
nulos_ano['Año'] = nulos_ano['Año'].astype(int)

temp_ano = df_v.groupby('source_file')['lst_celsius'].mean().reset_index()
temp_ano.columns = ['Año', 'LST media (C)']
temp_ano['LST media (C)'] = temp_ano['LST media (C)'].apply(lambda x: f"{x:.2f}")
temp_ano['Año'] = temp_ano['Año'].astype(int)

resumen_ano = nulos_ano.merge(temp_ano, on='Año')

temp_dpto = (df_v.groupby('ADM1_NAME')['lst_celsius']
             .agg(['mean','std']).sort_values('mean', ascending=False).reset_index())
temp_dpto.columns = ['Departamento', 'Media (C)', 'Std (C)']
temp_dpto['Media (C)'] = temp_dpto['Media (C)'].apply(lambda x: f"{x:.2f}")
temp_dpto['Std (C)']   = temp_dpto['Std (C)'].apply(lambda x: f"{x:.2f}")

print("Creando Word...")
doc = Document()
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(1.2)

heading(doc, 'Exploracion de Datos — Temperatura Superficial LST MODIS (GEE)', 1)
body(doc, 'Sistema de Alerta Temprana de Dengue · Grupo 11 · MAIA PDS · Universidad de los Andes')
body(doc, 'Fuente: Google Earth Engine (MODIS Land Surface Temperature) · Periodo: 2007–2024 · Granularidad: diaria por municipio')
doc.add_paragraph()

heading(doc, '1. Resumen general', 2)
res = pd.DataFrame({
    'Parametro': [
        'Total filas', 'Filas con datos validos', 'Municipios', 'Departamentos',
        'Granularidad', 'Periodo', 'Variable principal', 'Factor de escala MODIS',
        'Nulos promedio (nubosidad)', 'LST media Colombia', 'LST minima', 'LST maxima'
    ],
    'Valor': [
        '7,140,450', f"{len(df_v):,} ({len(df_v)/len(df)*100:.1f}%)",
        '1,086', '33', 'Diaria (un registro por municipio por dia)',
        '2007–2024 (18 anos)', 'lst_celsius — temperatura superficial del suelo (C)',
        'mean_raw x 0.02 - 273.15 = Celsius',
        f"{df['lst_celsius'].isna().mean()*100:.1f}%",
        f"{df_v['lst_celsius'].mean():.2f} C",
        f"{df_v['lst_celsius'].min():.2f} C",
        f"{df_v['lst_celsius'].max():.2f} C"
    ]
})
stats_table(doc, res)
body(doc, (
    'Los datos provienen del producto MODIS LST (MOD11A1) procesado en Google Earth Engine. '
    'Cada fila representa la temperatura media diaria de la superficie del suelo para un '
    'municipio colombiano. El alto porcentaje de nulos (59.7%) es inherente al sensor — '
    'la cobertura nubosa impide la lectura infrarroja, especialmente en la region Andina y '
    'el Pacifico. Los valores extremos negativos (-37°C) corresponden a areas de alta montana '
    '(picos nevados) y son fisicamente plausibles para superficies glaciares.'
))

heading(doc, '2. Datos faltantes por ano (nubosidad)', 2)
body(doc, (
    'La tasa de nulos es estable entre 57.5% y 63.4% a lo largo del periodo. '
    '2010 presenta el mayor porcentaje (63.4%), consistente con el fenomeno El Nino '
    'de ese ano que incremento la nubosidad en varias regiones. '
    '2018 tiene la menor tasa de nulos (57.5%), indicando mejores condiciones de cielo despejado.'
))
add_figure(doc, os.path.join(FIGS, '06_gee_nulos_por_ano.png'),
           'Figura 6. Porcentaje de datos faltantes por ano. La nubosidad es la principal causa de nulos en MODIS LST.')
stats_table(doc, resumen_ano)

heading(doc, '3. Estacionalidad — temperatura mensual', 2)
body(doc, (
    'Colombia presenta una estacionalidad termica moderada asociada a los periodos secos '
    '(diciembre–marzo y julio–agosto) y lluviosos (abril–junio y septiembre–noviembre). '
    'Los meses secos muestran temperaturas superficiales mas altas por mayor radiacion solar '
    'directa. El patron es consistente entre anos.'
))
add_figure(doc, os.path.join(FIGS, '07_gee_temp_mensual.png'),
           'Figura 7. Temperatura superficial mensual 2007–2024. Lineas azules: anos individuales. Roja: promedio historico.')

heading(doc, '4. Tendencia anual', 2)
body(doc, (
    'La temperatura superficial promedio nacional muestra variaciones interanuales '
    'asociadas a fenomenos climaticos (El Nino/La Nina). No se observa una tendencia '
    'de calentamiento estadisticamente clara en el periodo 2007–2024, aunque anos '
    'El Nino tienden a mostrar temperaturas mas altas.'
))
add_figure(doc, os.path.join(FIGS, '08_gee_temp_por_ano.png'),
           'Figura 8. Temperatura superficial promedio anual — Colombia.')

heading(doc, '5. Distribucion de temperatura', 2)
body(doc, (
    'La distribucion de LST es aproximadamente normal con media 25.17°C y desviacion '
    'estandar 5.62°C, reflejando la diversidad climatica de Colombia: desde municipios '
    'costeros y llaneros con temperaturas superiores a 40°C, hasta municipios de alta '
    'montana con valores negativos en superficies glaciares.'
))
add_figure(doc, os.path.join(FIGS, '09_gee_distribucion_lst.png'),
           'Figura 9. Distribucion de temperatura superficial (valores validos, 2007–2024).')

heading(doc, '6. Distribucion geografica por departamento', 2)
body(doc, (
    'Los departamentos de la Costa Caribe, Orinoquia y Amazonia presentan las '
    'temperaturas superficiales mas altas (>30°C), mientras que Boyaca, Narino y '
    'Cundinamarca registran las mas bajas por su mayor altitud media. Esta variabilidad '
    'geografica es un predictor relevante del riesgo de dengue.'
))
add_figure(doc, os.path.join(FIGS, '10_gee_temp_por_dpto.png'),
           'Figura 10. Temperatura superficial promedio por departamento.')
stats_table(doc, temp_dpto)

heading(doc, '7. Consideraciones para el modelado', 2)
body(doc, (
    'Para integrar los datos de LST con SIVIGILA, el campo de union es '
    'ADM2_CODE (GEE) = COD_MUN_O (SIVIGILA). Antes del modelado se requiere '
    'agregar la temperatura diaria a granularidad semanal (semana epidemiologica) '
    'e imputar los nulos — las estrategias recomendadas son interpolacion lineal '
    'temporal por municipio o imputacion por la mediana de la misma semana en '
    'anos no nulos. La temperatura es un predictor biologicamente justificado del '
    'ciclo reproductivo del vector Aedes aegypti (temperatura optima: 25–30°C).'
))

doc.save(OUT)
print(f"Guardado: {OUT}")
